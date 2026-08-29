# -*- coding: utf-8 -*-
import os
import shutil
import pytest
import docx
from fastapi.testclient import TestClient

from app.core.loader import DocumentLoader
from app.core.splitter import MarkdownHeaderSplitter
from app.core.bm25 import BM25Store
from app.core.reranker import RerankerClient
from app.main import app


@pytest.fixture(scope="session")
def temp_dir():
    test_dir = "./temp_test_data"
    os.makedirs(test_dir, exist_ok=True)
    yield test_dir
    if os.path.exists(test_dir):
        shutil.rmtree(test_dir)


def test_markdown_loader_and_splitter(temp_dir):
    md_content = """# 系统架构指南

本文档介绍系统的整体架构。

## 存储层

存储层负责保存所有持久化数据。

### ChromaDB 向量库
ChromaDB 是一个轻量级嵌入式向量数据库，支持余弦相似度检索。

### BM25 稀疏索引
BM25 用于进行精确的中文关键词倒排索引召回。

## 服务层
FastAPI 提供高性能异步接口支持。
"""
    md_file = os.path.join(temp_dir, "test_arch.md")
    with open(md_file, "w", encoding="utf-8") as f:
        f.write(md_content)

    # 1. 测试加载
    content, meta = DocumentLoader.load(md_file)
    assert "# 系统架构指南" in content
    assert meta["file_name"] == "test_arch.md"

    # 2. 测试切片与标题面包屑
    splitter = MarkdownHeaderSplitter(chunk_size=100, chunk_overlap=10)
    chunks = splitter.split_text(content, doc_metadata=meta)
    assert len(chunks) >= 3

    # 验证面包屑是否正确注入
    chroma_chunk = next((c for c in chunks if "ChromaDB 是一个轻量级" in c.content), None)
    assert chroma_chunk is not None
    assert "系统架构指南 > 存储层 > ChromaDB 向量库" in chroma_chunk.header_path
    assert "[上下文: 系统架构指南 > 存储层 > ChromaDB 向量库]" in chroma_chunk.content


def test_docx_loader_and_table(temp_dir):
    docx_file = os.path.join(temp_dir, "test_table.docx")
    doc = docx.Document()
    doc.add_heading("技术参数表", level=1)
    doc.add_paragraph("以下是服务器配置参数：")

    table = doc.add_table(rows=3, cols=2)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = '配置项'
    hdr_cells[1].text = '规格说明'
    
    r1 = table.rows[1].cells
    r1[0].text = 'CPU'
    r1[1].text = '16核 AMD EPYC'

    r2 = table.rows[2].cells
    r2[0].text = '内存'
    r2[1].text = '64GB DDR5 ECC'

    doc.save(docx_file)

    # 测试解析 DOCX 表格转 Markdown
    content, meta = DocumentLoader.load(docx_file)
    assert "# 技术参数表" in content
    assert "| 配置项 | 规格说明 |" in content
    assert "| CPU | 16核 AMD EPYC |" in content


def test_bm25_store(temp_dir):
    bm25 = BM25Store(persist_dir=os.path.join(temp_dir, "bm25"))
    kb_name = "test_kb"
    
    sample_chunks = [
        {
            "chunk_id": "chunk-1",
            "content": "[上下文: 产品说明] 本系统专为企业级知识库智能检索设计，支持流式输出。",
            "metadata": {"doc_name": "doc1.md"}
        },
        {
            "chunk_id": "chunk-2",
            "content": "[上下文: 数据库] ChromaDB 采用本地持久化方式保存向量维度数据。",
            "metadata": {"doc_name": "doc2.md"}
        }
    ]

    bm25.add_chunks(kb_name, sample_chunks)
    
    # 检索命中测试
    results = bm25.search(kb_name, "企业级知识库智能检索", top_k=2)
    assert len(results) >= 1
    assert results[0][0]["chunk_id"] == "chunk-1"
    assert results[0][1] > 0


def test_reranker_fallback():
    # 测试未配置 API Key 时的平滑回退
    client = RerankerClient(api_key="")
    candidates = [
        ({"chunk_id": "c1", "content": "第一篇"}, 0.9),
        ({"chunk_id": "c2", "content": "第二篇"}, 0.7),
        ({"chunk_id": "c3", "content": "第三篇"}, 0.5)
    ]
    reranked = client.rerank("测试查询", candidates, top_n=2)
    assert len(reranked) == 2
    assert reranked[0][0]["chunk_id"] == "c1"


def test_fastapi_health():
    client = TestClient(app)
    response = client.get("/health")
    assert response.status_code == 200
    data = response.json()
    assert data["status"] == "healthy"
    assert data["service"] == "RAG-GK Engine"


def test_fastapi_kb_lifecycle():
    client = TestClient(app)
    kb_name = "pytest_test_kb"

    # 1. 创建知识库
    create_res = client.post("/api/v1/kb/create", json={"kb_name": kb_name, "description": "Pytest KB"})
    assert create_res.status_code == 200
    assert create_res.json()["kb_name"] == kb_name

    # 2. 列出知识库
    list_res = client.get("/api/v1/kb/list")
    assert list_res.status_code == 200
    kbs = [item["kb_name"] for item in list_res.json()["knowledge_bases"]]
    assert kb_name in kbs

    # 3. 删除知识库
    del_res = client.delete(f"/api/v1/kb/{kb_name}")
    assert del_res.status_code == 200


def test_fastapi_non_stream_chat(monkeypatch):
    """非流式分支应调用同步生成器并返回标准 JSON。"""
    monkeypatch.setattr(
        "app.api.chat_router.retriever.retrieve",
        lambda **kwargs: []
    )
    monkeypatch.setattr(
        "app.api.chat_router.generator.generate_sync",
        lambda **kwargs: {
            "answer": "测试回答",
            "references": [],
            "usage": {
                "prompt_tokens": 1,
                "completion_tokens": 2,
                "total_tokens": 3,
            },
        }
    )

    client = TestClient(app)
    response = client.post(
        "/api/v1/chat/completions",
        json={
            "kb_name": "test_kb",
            "query": "测试问题",
            "stream": False,
        },
    )

    assert response.status_code == 200
    assert response.json() == {
        "answer": "测试回答",
        "references": [],
        "usage": {
            "prompt_tokens": 1,
            "completion_tokens": 2,
            "total_tokens": 3,
        },
    }
